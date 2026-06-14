"""Part 4: Final sections 8-11 + simulation animation video + finish document"""
import os, sys, site, glob

user_site = site.getusersitepackages()
if user_site not in sys.path:
    sys.path.insert(0, user_site)

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import matplotlib.animation as animation
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE     = "/home/jatin/ros2_assign_2_10_robo_modified_tonew/turtle_to_quadcopter final with modifications"
IMG_BASE = "/mnt/c/Users/HP/.gemini/antigravity/brain/52bd95c7-02fd-43b7-aea1-8cf37fce8c85"
OUT_DOCX = os.path.join(BASE, "Quadcopter_MultiAgent_Report.docx")

def find_img(prefix):
    hits = glob.glob(os.path.join(IMG_BASE, f"{prefix}_*.png"))
    return hits[0] if hits else None

GRID_IMG = find_img("grid_map_diagram")

# ── Simulation Screenshot (static matplotlib render) ─────────────────────────
def make_sim_snapshot():
    """Renders a top-down view of the 10x8 grid with 4 drones at sample positions."""
    fig, ax = plt.subplots(figsize=(10,7), facecolor="#0d1117")
    ax.set_facecolor("#0f1923")
    ax.set_xlim(-1, 19); ax.set_ylim(-1, 15)
    ax.set_aspect("equal"); ax.axis("off")
    ax.set_title("Simulation State Snapshot — 4 Drones in FLYING Mode",
                 color="white", fontsize=13, fontweight="bold", pad=10)

    rows = list("ABCDEFGH")
    obstacles = {"B4","B8","C2","C6","D5","D9","E3","E7","F2","F6","G4","G9"}

    # Draw grid cells
    for ri, row in enumerate(rows):
        for ci in range(1, 11):
            name = f"{row}{ci}"
            x, y = (ci-1)*2, ri*2
            is_obs = name in obstacles
            color = "#8b1a1a" if is_obs else "#1a2535"
            ec    = "#cc2200" if is_obs else "#2a3f5f"
            lw    = 2.0 if is_obs else 0.8
            z     = 4 if is_obs else 1
            rect  = plt.Rectangle((x-0.9, y-0.9), 1.8, 1.8,
                                   facecolor=color, edgecolor=ec, linewidth=lw, zorder=z)
            ax.add_patch(rect)
            fc = "#ff4444" if is_obs else "#4a7fa5"
            ax.text(x, y, name, ha="center", va="center", color=fc,
                    fontsize=6, fontweight="bold", zorder=2)

    # Draw drone paths (simplified)
    # All paths pre-verified to avoid obstacle cells:
    # Obstacles(col,row->x,y): B4(6,2),B8(14,2),C2(2,4),C6(10,4),D5(8,6),D9(16,6)
    #                          E3(4,8),E7(12,8),F2(2,10),F6(10,10),G4(6,12),G9(16,12)
    paths = {
        "robot1": [(0,0),(2,0),(4,0),(6,0),(8,0),(10,0),(12,0),(14,0),(16,0),(18,0)],
        "robot2": [(0,0),(0,2),(0,4),(0,6),(0,8),(0,10),(0,12),(0,14),(2,14),(4,14)],
        "robot3": [(18,14),(16,14),(14,14),(12,14),(10,14),(8,14),(8,12),(8,10),(8,8),(6,8),(6,6),(4,6)],
        "robot4": [(18,0),(18,2),(18,4),(16,4),(14,4),(12,4),(12,6),(14,6),(14,8),(14,10)],
    }
    drone_colors = {"robot1":"#e74c3c","robot2":"#3498db","robot3":"#2ecc71","robot4":"#9b59b6"}
    drone_pos    = {"robot1":(10,0),"robot2":(0,8),"robot3":(8,10),"robot4":(14,6)}

    for rid, path in paths.items():
        xs = [p[0] for p in path]; ys = [p[1] for p in path]
        ax.plot(xs, ys, color=drone_colors[rid], linewidth=1.5, alpha=0.5, linestyle="--", zorder=3)

    for rid, (dx,dy) in drone_pos.items():
        col = drone_colors[rid]
        # Comm disc
        disc = plt.Circle((dx, dy), 1.8, color=col, alpha=0.12, zorder=3)
        ax.add_patch(disc)
        # Body
        body = plt.Rectangle((dx-0.2, dy-0.2), 0.4, 0.4, color=col, zorder=5)
        ax.add_patch(body)
        # Arms (X configuration)
        for dx2, dy2 in [(0.4, 0.4), (-0.4, 0.4), (-0.4, -0.4), (0.4, -0.4)]:
            ax.plot([dx, dx+dx2], [dy, dy+dy2], color="#888", linewidth=2.5, zorder=4)
            prop = plt.Circle((dx+dx2, dy+dy2), 0.2, color=col, alpha=0.7, zorder=5)
            ax.add_patch(prop)
        # Label
        ax.text(dx, dy+0.7, rid, ha="center", va="bottom", color="white",
                fontsize=7.5, fontweight="bold", zorder=6,
                bbox=dict(facecolor="#00000088", edgecolor="none", pad=1))

    ax.text(9, -0.6, "Columns: 1  2  3  4  5  6  7  8  9  10",
            ha="center", color="#7090b0", fontsize=8)

    path_out = os.path.join(BASE, "sim_snapshot.png")
    plt.tight_layout()
    plt.savefig(path_out, dpi=150, bbox_inches="tight", facecolor="#0d1117")
    plt.close()
    return path_out

# ── Simulation Animation GIF ──────────────────────────────────────────────────
def make_sim_animation():
    """Creates an animated GIF of drones moving on the grid."""
    fig, ax = plt.subplots(figsize=(8,6), facecolor="#0d1117")
    ax.set_facecolor("#0f1923"); ax.set_xlim(-1,19); ax.set_ylim(-1,15)
    ax.set_aspect("equal"); ax.axis("off")

    rows_l = list("ABCDEFGH")
    obstacles = {"B4","B8","C2","C6","D5","D9","E3","E7","F2","F6","G4","G9"}

    for ri, row in enumerate(rows_l):
        for ci in range(1,11):
            name = f"{row}{ci}"
            x, y = (ci-1)*2, ri*2
            is_obs = name in obstacles
            color = "#8b1a1a" if is_obs else "#1a2535"
            ec    = "#cc2200" if is_obs else "#2a3f5f"
            lw    = 2.0 if is_obs else 0.8
            z     = 5 if is_obs else 1   # obstacles always on top of comm discs
            rect  = plt.Rectangle((x-0.9,y-0.9),1.8,1.8,facecolor=color,edgecolor=ec,linewidth=lw,zorder=z)
            ax.add_patch(rect)
            fc = "#ff6666" if is_obs else "#2a4a6a"
            tz = 6 if is_obs else 2
            ax.text(x,y,name,ha="center",va="center",color=fc,fontsize=5.5,fontweight="bold",zorder=tz)

    DRONE_COLS = ["#e74c3c","#3498db","#2ecc71","#9b59b6"]
    NAMES      = ["robot1","robot2","robot3","robot4"]
    # Verified obstacle-free waypoints (x=(col-1)*2, y=row_index*2)
    WAYPOINTS  = [
        # robot1 (red):   Row A left→right, then col 10 down
        [(0,0),(2,0),(4,0),(6,0),(8,0),(10,0),(12,0),(14,0),(16,0),(18,0),(18,2),(18,4),(18,6),(18,8)],
        # robot2 (blue):  Col 1 down, then row H right
        [(0,0),(0,2),(0,4),(0,6),(0,8),(0,10),(0,12),(0,14),(2,14),(4,14),(6,14),(8,14),(10,14),(12,14)],
        # robot3 (green): Row H right→left, then interior avoiding G4(6,12)
        [(18,14),(16,14),(14,14),(12,14),(10,14),(8,14),(8,12),(8,10),(8,8),(6,8),(6,6),(4,6),(4,4),(4,2)],
        # robot4 (purple): Col 10 up, interior path avoiding D9(16,6)
        [(18,0),(18,2),(18,4),(16,4),(14,4),(12,4),(12,6),(14,6),(14,8),(14,10),(14,12),(14,14)],
    ]

    drone_body_artists = []
    drone_arm_artists = []
    drone_prop_artists = []
    label_artists = []
    disc_artists  = []
    arm_offsets = [(0.4, 0.4), (-0.4, 0.4), (-0.4, -0.4), (0.4, -0.4)]

    for i, (col, name) in enumerate(zip(DRONE_COLS, NAMES)):
        disc = plt.Circle((0,0), 1.8, color=col, alpha=0.12, zorder=3)
        ax.add_patch(disc); disc_artists.append(disc)
        body = plt.Rectangle((-0.2,-0.2), 0.4, 0.4, color=col, zorder=5)
        ax.add_patch(body); drone_body_artists.append(body)
        
        arms = []
        props = []
        for dx2, dy2 in arm_offsets:
            arm, = ax.plot([0, dx2], [0, dy2], color="#888", linewidth=2.5, zorder=4)
            arms.append(arm)
            prop = plt.Circle((dx2, dy2), 0.2, color=col, alpha=0.7, zorder=5)
            ax.add_patch(prop)
            props.append(prop)
            
        drone_arm_artists.append(arms)
        drone_prop_artists.append(props)

        lbl  = ax.text(0,0.8,name,ha="center",va="bottom",color="white",fontsize=6.5,
                       fontweight="bold",zorder=6,bbox=dict(facecolor="#00000088",edgecolor="none",pad=1))
        label_artists.append(lbl)

    title_txt = ax.set_title("Quadcopter Multi-Agent Simulation — Step 0",
                              color="white", fontsize=11, fontweight="bold")

    NFRAMES = 40

    def update(frame):
        returned_artists = [title_txt]
        for i, (col, name, waypts) in enumerate(zip(DRONE_COLS, NAMES, WAYPOINTS)):
            idx = int(frame * len(waypts) / NFRAMES) % len(waypts)
            x, y = waypts[idx]
            drone_body_artists[i].set_xy((x-0.2, y-0.2))
            disc_artists[i].center = (x, y)
            label_artists[i].set_position((x, y+0.8))
            
            for j, (dx2, dy2) in enumerate(arm_offsets):
                drone_arm_artists[i][j].set_data([x, x+dx2], [y, y+dy2])
                drone_prop_artists[i][j].center = (x+dx2, y+dy2)
            
            returned_artists.extend([drone_body_artists[i], disc_artists[i], label_artists[i]])
            returned_artists.extend(drone_arm_artists[i])
            returned_artists.extend(drone_prop_artists[i])

        title_txt.set_text(f"Quadcopter Multi-Agent Simulation — Tick {frame}")
        return returned_artists

    ani = animation.FuncAnimation(fig, update, frames=NFRAMES, interval=150, blit=False)
    gif_path = os.path.join(BASE, "sim_animation.gif")
    ani.save(gif_path, writer="pillow", fps=8, dpi=100)
    plt.close()
    print(f"Animation saved: {gif_path}")
    return gif_path

print("Generating simulation snapshot..."); snap_img = make_sim_snapshot()
print("Generating simulation animation GIF..."); gif_path = make_sim_animation()

# ── Word helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell,hx):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
    shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hx); tcPr.append(shd)

def add_heading(doc,text,level=1,color="1E3A5F"):
    h=doc.add_heading(text,level=level); h.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in h.runs: r.font.color.rgb=RGBColor.from_string(color); r.font.bold=True

def add_para(doc,text,size=11,bold=False,color=None):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    r=p.add_run(text); r.font.size=Pt(size); r.font.bold=bold
    if color: r.font.color.rgb=RGBColor.from_string(color)

def add_code(doc,text):
    p=doc.add_paragraph(); p.style=doc.styles["No Spacing"]
    r=p.add_run(text); r.font.name="Courier New"; r.font.size=Pt(8.5)
    r.font.color.rgb=RGBColor(0x1A,0x5C,0x2A)

def add_img(doc,path,width=Inches(5.5),caption=None):
    if path and os.path.exists(path):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path,width=width)
        if caption:
            cp=doc.add_paragraph(caption); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in cp.runs: r.font.size=Pt(9); r.font.italic=True; r.font.color.rgb=RGBColor(100,100,100)
    else: doc.add_paragraph(f"[Image not found: {path}]")

def add_table(doc,headers,rows,hbg="1E3A5F",alt="EAF0FB"):
    t=doc.add_table(rows=1+len(rows),cols=len(headers))
    t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hr=t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.text=h; set_cell_bg(c,hbg)
        for p in c.paragraphs:
            for r in p.runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for ri,rd in enumerate(rows):
        row=t.rows[ri+1]; bg=alt if ri%2==0 else "FFFFFF"
        for ci,val in enumerate(rd):
            c=row.cells[ci]; c.text=val; set_cell_bg(c,bg)
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)

doc = Document(OUT_DOCX)

# ── SECTION 8: ROS TOPICS & SERVICES ─────────────────────────────────────────
add_heading(doc,"8. ROS Topics, Services & Parameters",level=1)
add_heading(doc,"Topics",level=2,color="2E5FA0")
add_table(doc,["Topic","Msg Type","Publisher","Subscribers","Description"],[
    ("/robot_cell_registry","String","All robot_nodes","All robot_nodes + graph_vis","Drone cell reservation registry"),
    ("/robotN/pose","Pose","robot_node N","All robot_nodes + prox_mon","3D position + quaternion"),
    ("/robotN/markers","MarkerArray","robot_node N","RViz2","Full drone visual"),
    ("/graph_markers","MarkerArray","graph_visualizer","RViz2","Grid + obstacles + labels"),
    ("/dynamic_obstacles","String","graph_visualizer","All robot_nodes","New obstacle cell list"),
])
doc.add_paragraph("")
add_heading(doc,"Services",level=2,color="2E5FA0")
add_table(doc,["Service","Type","Server","Client","Trigger Condition"],[
    ("/robotN/communicate","Trigger","robot_node N","proximity_monitor","Two drones within threshold distance"),
])
doc.add_paragraph("")
add_heading(doc,"Parameters by Node",level=2,color="2E5FA0")
add_table(doc,["Node","Parameter","Type","Default","Description"],[
    ("robot_node","robot_id","string","robot1","Unique drone ID"),
    ("robot_node","color","string","red","Body and propeller colour"),
    ("robot_node","path","string","A1,A2,A3","Waypoint list (comma-separated)"),
    ("robot_node","speed","float","0.4","Virtual rabbit speed (m/s)"),
    ("robot_node","radius","float","2.0","Comm-range disc radius (m)"),
    ("proximity_monitor","threshold","float","10.0","Proximity detection radius (m)"),
    ("proximity_monitor","robot_ids","string","—","Comma-separated drone IDs"),
    ("graph_visualizer","spawn_interval","float","8.0","Obstacle reshuffle interval (s)"),
])
doc.add_page_break()

# ── SECTION 9: RUNNING THE PROJECT ───────────────────────────────────────────
add_heading(doc,"9. Running the Project — All Modes",level=1)

# Simulation visuals
add_img(doc, snap_img, width=Inches(5.8),
        caption="Figure 12 — Top-down simulation snapshot: 4 drones in FLYING state with dashed path trails")
doc.add_paragraph("")
add_img(doc, gif_path, width=Inches(5.0),
        caption="Figure 13 — Animated simulation GIF: drones traversing the 10×8 grid (embedded frame)")
doc.add_paragraph("")

add_heading(doc,"Mode 1: Interactive Runner (Recommended)",level=2,color="2E5FA0")
add_code(doc,"ros2 run robot_proximity interactive_runner")
add_para(doc,"Prompts in order:")
add_table(doc,["Step","Prompt","Options"],[
    ("1","Display grid with [XX] obstacles","—"),
    ("2","Simulation mode?","1 = AUTO, 2 = MANUAL"),
    ("3","Obstacle spawn interval (seconds)?","0 = disable, e.g. 8"),
    ("4a (AUTO)","Number of robots? (1-10)","Selects first N pre-stored routes"),
    ("4b (MANUAL)","Enter waypoints per robot","e.g. A1,B1,C1,D1,H1"),
    ("5","Generates temp_interactive.launch.py","Launches via subprocess.run"),
])
doc.add_paragraph("")

add_heading(doc,"Mode 2: Static Launch File",level=2,color="2E5FA0")
add_code(doc,"ros2 launch robot_proximity simulation.launch.py")
add_para(doc,"Immediately launches 4 pre-configured robots. Dynamic obstacles use the default 8-second interval.")
doc.add_paragraph("")

add_heading(doc,"Mode 3: Individual Node Launch (Debug)",level=2,color="2E5FA0")
add_code(doc,
"""# Terminal 1 — Single drone
source /opt/ros/humble/setup.bash && source install/setup.bash
ros2 run robot_proximity robot_node --ros-args \\
  -p robot_id:=robot1 -p color:=red \\
  -p path:=A1,B1,C1,D1,H1 -p speed:=0.4 -p radius:=2.0

# Terminal 2 — Grid visualizer
ros2 run robot_proximity graph_visualizer --ros-args -p spawn_interval:=6.0

# Terminal 3 — Proximity monitor
ros2 run robot_proximity proximity_monitor --ros-args \\
  -p robot_ids:=robot1 -p threshold:=5.0

# Terminal 4 — RViz2
rviz2 -d src/robot_proximity/config/simulation.rviz""")
doc.add_paragraph("")

add_heading(doc,"Mode 4: Path Diagnostics (no ROS required)",level=2,color="2E5FA0")
add_code(doc,'python3 src/robot_proximity/robot_proximity/diag_test.py')
add_para(doc,"Tests three hard-coded paths offline against the GraphManager. No ROS installation needed.")
doc.add_page_break()

# ── SECTION 10: BUILD INSTRUCTIONS ───────────────────────────────────────────
add_heading(doc,"10. Build Instructions",level=1)
add_table(doc,["Step","Command","Purpose"],[
    ("1","source /opt/ros/humble/setup.bash","Load ROS 2 Humble environment variables"),
    ("2",'cd "turtle_to_quadcopter final with modifications"',"Navigate to workspace root"),
    ("3","colcon build --packages-select robot_proximity","Build only the robot_proximity package"),
    ("4","source install/setup.bash","Source the local install overlay"),
    ("5","ros2 run robot_proximity interactive_runner","Launch the interactive simulation"),
])
doc.add_paragraph("")
add_para(doc,"⚠  After any source code change, repeat Steps 3 and 4 before relaunching.",
         bold=True, color="CC4400")
add_para(doc,"Full build command block:",bold=True)
add_code(doc,
"""source /opt/ros/humble/setup.bash
cd "/home/jatin/ros2_assign_2_10_robo_modified_tonew/turtle_to_quadcopter final with modifications"
colcon build --packages-select robot_proximity
source install/setup.bash
ros2 run robot_proximity interactive_runner""")
doc.add_page_break()

# ── SECTION 11: GRID REFERENCE ───────────────────────────────────────────────
add_heading(doc,"11. Grid Reference",level=1)
add_img(doc, GRID_IMG, width=Inches(5.5),
        caption="Figure 14 — Full 10×8 grid reference: [XXX] = obstacle, walkable cells in blue-grey")
doc.add_paragraph("")
add_table(doc,["Property","Value"],[
    ("Dimensions","10 columns × 8 rows (A–H rows, 1–10 columns)"),
    ("Cell spacing","2.0 metres between adjacent cell centres"),
    ("Total cells","80 (A1 through H10)"),
    ("Static obstacles","12 cells: B4, B8, C2, C6, D5, D9, E3, E7, F2, F6, G4, G9"),
    ("Max dynamic obstacles","15 (randomly sampled unoccupied cells)"),
    ("Walkable (static only)","68 cells"),
    ("Grid extents","X: 0–18 m, Y: 0–14 m (Cartesian)"),
    ("Corner A1","(0.0, 0.0)"),
    ("Corner H10","(18.0, 14.0)"),
])
doc.add_paragraph("")

# ── COLOPHON / SIGNATURE ──────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("─── End of Report ───")
r.font.size=Pt(12); r.font.color.rgb=RGBColor(0x88,0x99,0xAA); r.font.italic=True
doc.add_paragraph("")
for line in [
    "robot_proximity  v0.0.0",
    "ROS 2 Humble Hawksbill  |  Python 3.10  |  Apache-2.0",
    "Maintainer: Jatin",
]:
    lp = doc.add_paragraph(); lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = lp.add_run(line); lr.font.size=Pt(10); lr.font.color.rgb=RGBColor(0x66,0x77,0x88)

doc.save(OUT_DOCX)
print(f"\n✅ FINAL REPORT COMPLETE → {OUT_DOCX}")
print("   Sections added: 8 (Topics/Services), 9 (Running), 10 (Build), 11 (Grid Ref)")
print(f"   Simulation snapshot : {snap_img}")
print(f"   Animation GIF       : {gif_path}")
print("\nAll 4 parts done! Open the .docx file to see the full report.")
