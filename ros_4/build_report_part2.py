"""Part 2: Append Sections 3-5 to existing Word doc"""
import os, sys, site, glob

user_site = site.getusersitepackages()
if user_site not in sys.path:
    sys.path.insert(0, user_site)
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE     = "/home/jatin/ros2_assign_2_10_robo_modified_tonew/turtle_to_quadcopter final with modifications"
IMG_BASE = "/mnt/c/Users/HP/.gemini/antigravity/brain/52bd95c7-02fd-43b7-aea1-8cf37fce8c85"
OUT_DOCX = os.path.join(BASE, "Quadcopter_MultiAgent_Report.docx")

def find_img(prefix):
    hits = glob.glob(os.path.join(IMG_BASE, f"{prefix}_*.png"))
    return hits[0] if hits else None

def find_chart(name):
    p = os.path.join(BASE, name)
    return p if os.path.exists(p) else None

ROS_IMG  = find_img("ros_architecture_diagram")
DRONE_IMG= find_img("drone_architecture_diagram")
GRID_IMG = find_img("grid_map_diagram")
TIMER_IMG= find_chart("chart_timer_arch.png")
PIE_IMG  = find_chart("chart_grid_pie.png")

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="1E3A5F"):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color); run.font.bold = True

def add_para(doc, text, size=11, bold=False, color=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text); r.font.size=Pt(size); r.font.bold=bold
    if color: r.font.color.rgb = RGBColor.from_string(color)

def add_code(doc, text):
    p = doc.add_paragraph(); p.style = doc.styles["No Spacing"]
    r = p.add_run(text); r.font.name="Courier New"; r.font.size=Pt(8.5)
    r.font.color.rgb = RGBColor(0x1A,0x5C,0x2A)

def add_img(doc, path, width=Inches(5.5), caption=None):
    if path and os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=width)
        if caption:
            cp = doc.add_paragraph(caption); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cp.runs:
                r.font.size=Pt(9); r.font.italic=True; r.font.color.rgb=RGBColor(100,100,100)
    else:
        doc.add_paragraph(f"[Image not found: {path}]")

def add_table(doc, headers, rows, hbg="1E3A5F", alt="EAF0FB"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.text=h; set_cell_bg(c,hbg)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for ri,rd in enumerate(rows):
        row=t.rows[ri+1]; bg=alt if ri%2==0 else "FFFFFF"
        for ci,val in enumerate(rd):
            c=row.cells[ci]; c.text=val; set_cell_bg(c,bg)
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)

doc = Document(OUT_DOCX)

# ── SECTION 3: ROS 2 PACKAGE CONFIGURATION ───────────────────────────────────
add_heading(doc,"3. ROS 2 Package Configuration",level=1)
add_para(doc,"The package.xml declares the package name robot_proximity, build type ament_python, and core ROS 2 dependencies:")
add_table(doc,["Dependency","What It Provides"],[
    ("rclpy","Python ROS 2 client: Node, Publisher, Subscriber, Timer, Service APIs"),
    ("geometry_msgs","Pose message for broadcasting drone 3D position + quaternion orientation"),
    ("visualization_msgs","Marker and MarkerArray for rendering all 3D objects in RViz"),
    ("std_msgs","String message for the cell registry topic"),
    ("std_srvs","Trigger service definition for inter-drone communication RPC"),
])
doc.add_paragraph("")
add_heading(doc,"setup.py Entry Points",level=2,color="2E5FA0")
add_para(doc,"Four console-script entry points registered via colcon:")
add_code(doc,
"""'console_scripts': [
    'robot_node        = robot_proximity.robot_node:main',
    'proximity_monitor = robot_proximity.proximity_monitor:main',
    'graph_visualizer  = robot_proximity.graph_visualizer:main',
    'interactive_runner= robot_proximity.interactive_runner:main',
]""")
doc.add_page_break()

# ── SECTION 4: SYSTEM ARCHITECTURE ───────────────────────────────────────────
add_heading(doc,"4. System Architecture & Node Communication",level=1)
add_para(doc,(
    "The system is completely decentralized — there is no master controller. All coordination "
    "flows through the ROS 2 middleware. Every robot_node independently subscribes to all other "
    "drones' pose and cell data and makes its own navigation decisions."
))
doc.add_paragraph("")
add_img(doc, ROS_IMG, width=Inches(5.8),
        caption="Figure 2 — ROS 2 node communication topology (publish/subscribe + service mesh)")
doc.add_paragraph("")
add_para(doc,"Communication flow summary:",bold=True)
add_table(doc,["Channel","Publisher","Subscribers","Type"],[
    ("/robot_cell_registry","All robot_nodes","All robot_nodes + graph_visualizer","String"),
    ("/robotN/pose","robot_node N","All robot_nodes + proximity_monitor","Pose"),
    ("/robotN/markers","robot_node N","RViz2","MarkerArray"),
    ("/graph_markers","graph_visualizer","RViz2","MarkerArray"),
    ("/dynamic_obstacles","graph_visualizer","All robot_nodes","String"),
    ("/robotN/communicate","—","robot_node N (server)","Trigger Service"),
])
doc.add_page_break()

# ── SECTION 5: COMPONENT DEEP DIVE ───────────────────────────────────────────
add_heading(doc,"5. Component Deep Dive",level=1)

# 5.1 graph_manager
add_heading(doc,"5.1  graph_manager.py — Grid and BFS Engine",level=2,color="2E5FA0")
add_para(doc,(
    "The mathematical and topological foundation of the entire project. Instantiated independently "
    "by robot_node, graph_visualizer, and interactive_runner."
))
add_para(doc,"Grid Construction:",bold=True)
add_para(doc,(
    "A 10-column × 8-row matrix (A1 to H10, 80 nodes total) with 2.0 m cell spacing. "
    "12 static obstacle cells block traversal: B4, B8, C2, C6, D5, D9, E3, E7, F2, F6, G4, G9."
))
doc.add_paragraph("")
add_img(doc, GRID_IMG, width=Inches(5.5),
        caption="Figure 3 — 10×8 simulation grid with static obstacle cells (red blocks)")
doc.add_paragraph("")
add_img(doc, PIE_IMG, width=Inches(4.5),
        caption="Figure 4 — Grid cell distribution: static obstacles / dynamic obstacles / walkable")
doc.add_paragraph("")
add_para(doc,"Key API Methods:",bold=True)
add_table(doc,["Method","Returns","Description"],[
    ("get_coords(node)","(float, float)","Cartesian XY coordinates of a node"),
    ("find_path(start, end)","list or None","BFS shortest path from start to end"),
    ("find_path_excluding(start,end,excl)","list or None","BFS avoiding a set of nodes (deadlock rerouting)"),
    ("resolve_path(waypoints)","list","Expands sparse waypoints into full edge-by-edge path"),
    ("update_obstacles(new_obs)","None","Replaces obstacle set and rebuilds edges + adjacency"),
    ("is_obstacle(node)","bool","True if the node is currently blocked"),
    ("interpolate(start,end,alpha)","(float,float)","Linear interpolation between two nodes"),
])
doc.add_paragraph("")

# 5.2 robot_node
add_heading(doc,"5.2  robot_node.py — Physics Agent & Flight Controller",level=2,color="2E5FA0")
add_para(doc,(
    "The core agent. One instance runs per drone. Manages 3D physics, flight states, "
    "all collision avoidance logic, and produces all RViz markers for its drone."
))
add_para(doc,"ROS 2 Parameters:",bold=True)
add_table(doc,["Parameter","Type","Default","Description"],[
    ("robot_id","string","robot1","Unique drone identifier"),
    ("color","string","red","Drone body and propeller colour"),
    ("path","string","A1,A2,A3","Comma-separated waypoint list"),
    ("speed","float","0.4","Virtual rabbit speed in m/s"),
    ("radius","float","2.0","Communication range disc radius in metres"),
])
doc.add_paragraph("")
add_para(doc,"Timer Architecture:",bold=True)
add_img(doc, TIMER_IMG, width=Inches(5.5),
        caption="Figure 5 — Dual-timer architecture: 50 Hz physics loop + 3 Hz marker refresh")
doc.add_paragraph("")
add_para(doc,"Drone Marker Composition (9 primitives per drone):",bold=True)
add_table(doc,["Marker ID","Type","Description"],[
    ("0","CUBE","Central fuselage (0.5×0.5×0.16 m), dimmed when waiting"),
    ("1","CYLINDER","Fore-aft arm (1.1 m length), dark grey, tilts with body"),
    ("2","CYLINDER","Left-right lateral arm, perpendicular to arm 1"),
    ("11-14","CYLINDER ×4","Motor housings at the four arm tips"),
    ("3-6","CYLINDER ×4","Propeller discs, coloured, counter-rotating pairs"),
    ("7","CYLINDER","Translucent comm-range disc on the ground"),
    ("8","TEXT_VIEW_FACING","Name label showing state: CRUISE/TAKEOFF/LANDING/YIELD/WAITING"),
])
doc.add_paragraph("")
add_img(doc, DRONE_IMG, width=Inches(5.0),
        caption="Figure 6 — Quadcopter drone structure: arms, motor housings, propellers, comm disc")
doc.add_page_break()

# 5.3 graph_visualizer
add_heading(doc,"5.3  graph_visualizer.py — Map Renderer & Obstacle Spawner",level=2,color="2E5FA0")
add_para(doc,(
    "Dual-purpose node: renders the full 10×8 grid in RViz and periodically randomises obstacles."
))
add_table(doc,["Publisher","Type","Rate","Description"],[
    ("/graph_markers","MarkerArray","1 Hz","Complete grid: cells, lines, labels, obstacles"),
    ("/dynamic_obstacles","String","user-defined","Comma-separated new obstacle node names"),
])
doc.add_paragraph("")
add_para(doc,"spawn_obstacles() Logic:",bold=True)
add_code(doc,
"""def spawn_obstacles(self):
    occupied = set()
    for cells in self.robot_cells.values():
        occupied.update(cells)            # Never block occupied cells
    possible = [n for n in self.gm.nodes if n not in occupied]
    new_obs  = random.sample(possible, min(15, len(possible)))
    msg.data = ",".join(new_obs)
    self.obs_pub.publish(msg)             # Broadcast to all robot_nodes
    self.gm.update_obstacles(new_obs)
    self.publish_graph()                  # Re-render RViz immediately""")
doc.add_paragraph("")

# 5.4 proximity_monitor
add_heading(doc,"5.4  proximity_monitor.py — Inter-Drone Observer",level=2,color="2E5FA0")
add_para(doc,(
    "Passive background observer. Tracks inter-drone distances and triggers "
    "ROS 2 Trigger service calls when drones come within range. Runs at 5 Hz."
))
add_table(doc,["Parameter","Default","Description"],[
    ("threshold","10.0","Distance in metres for proximity detection"),
    ("robot_ids","robot1,robot2,...","Comma-separated drone IDs to monitor"),
])
doc.add_paragraph("")
add_para(doc,"ASCII proximity log format example:",bold=True)
add_code(doc,
"""+------------------------------------------------------------+
| PROXIMITY DETECTED   (Interaction #7)                      |
|------------------------------------------------------------|
| Robot A :  robot1                                          |
| Robot B :  robot3                                          |
| Distance:  3.82 units  (threshold 10.0)                    |
+------------------------------------------------------------+""")
doc.add_paragraph("")

# 5.5 interactive_runner
add_heading(doc,"5.5  interactive_runner.py — Simulation Orchestrator",level=2,color="2E5FA0")
add_para(doc,(
    "Primary user-facing entry point. A pure Python terminal UI that gathers configuration, "
    "validates paths, generates a .launch.py file dynamically, then launches it via subprocess."
))
add_para(doc,"Startup Sequence:",bold=True)
add_table(doc,["Step","Action"],[
    ("1","Print 10×8 grid map showing [XX] obstacle cells"),
    ("2","Prompt: simulation mode — 1=AUTO or 2=MANUAL"),
    ("3","Prompt: obstacle spawn interval in seconds (0=disable)"),
    ("4","AUTO: prompt robot count (1-10) | MANUAL: enter waypoints per robot"),
    ("5","Validate all paths via gm.resolve_path()"),
    ("6","Generate temp_interactive.launch.py with all nodes configured"),
    ("7","Execute: subprocess.run(['ros2','launch', temp_launch_path])"),
])
doc.add_paragraph("")
add_para(doc,"AUTO_PATHS — 10 Pre-Stored Routes:",bold=True)
add_table(doc,["Robot","Colour","Route Description","Speed"],[
    ("robot1","red","Full top row A1→A10, then down right column to H10","0.8"),
    ("robot2","blue","Down left column A1→H1, then across bottom to H10","0.8"),
    ("robot3","green","S-curve through centre with obstacle auto-rerouting","0.8"),
    ("robot4","purple","Right-to-left top row A10→A1, then down left column","0.8"),
    ("robot5","orange","Column 5 full vertical traverse A5→H5","0.8"),
    ("robot6","cyan","Bottom row H10→H1, then up left column to A1","0.8"),
    ("robot7","magenta","Full outer perimeter clockwise loop","0.6"),
    ("robot8","yellow","Interior zigzag diagonal path","0.8"),
    ("robot9","teal","Column 6 traverse with horizontal cross at F row","0.8"),
    ("robot10","white","Inner rectangular loop C3 to G7","0.8"),
])
doc.add_page_break()

doc.save(OUT_DOCX)
print(f"✅ Part 2 SAVED → {OUT_DOCX}")
print("   Sections added: 3 (ROS Config), 4 (Architecture), 5 (Components)")
