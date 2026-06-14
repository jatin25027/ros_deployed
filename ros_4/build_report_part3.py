"""Part 3: Append Sections 6-7 (Algorithms + State Machine) to Word doc"""
import os, sys, site, glob

user_site = site.getusersitepackages()
if user_site not in sys.path:
    sys.path.insert(0, user_site)

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE     = "/home/jatin/ros2_assign_2_10_robo_modified_tonew/turtle_to_quadcopter final with modifications"
IMG_BASE = "/mnt/c/Users/HP/.gemini/antigravity/brain/52bd95c7-02fd-43b7-aea1-8cf37fce8c85"
OUT_DOCX = os.path.join(BASE, "Quadcopter_MultiAgent_Report.docx")

def find_img(prefix):
    hits = glob.glob(os.path.join(IMG_BASE, f"{prefix}_*.png"))
    return hits[0] if hits else None

FSM_IMG = find_img("flight_state_machine")
BFS_IMG = find_img("bfs_flowchart")

# ── Generate PD Controller Diagram ───────────────────────────────────────────
def make_pd_diagram():
    fig, ax = plt.subplots(figsize=(10,5), facecolor="#0d1117")
    ax.set_facecolor("#0d1117"); ax.set_xlim(0,10); ax.set_ylim(0,6); ax.axis("off")
    ax.set_title("PD Controller — Virtual Rabbit Tracking", color="white", fontsize=13, fontweight="bold")

    boxes = [
        (0.2,3.5,2.0,1.2,"#1565c0","Reference\n(Rabbit Pos)"),
        (2.8,3.5,1.8,1.2,"#6a1b9a","Summing\nJunction"),
        (5.2,3.5,2.0,1.2,"#00695c","PD Controller\nKp=3.5 Kd=2.0"),
        (7.8,3.5,1.8,1.2,"#b71c1c","Drone Physics\n(Euler Integ.)"),
        (7.8,1.5,1.8,1.2,"#37474f","Output\nDrone Pos"),
    ]
    for (x,y,w,h,col,txt) in boxes:
        rect = FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.1",
                              facecolor=col,edgecolor="white",linewidth=1.5,alpha=0.9)
        ax.add_patch(rect)
        ax.text(x+w/2,y+h/2,txt,ha="center",va="center",color="white",fontsize=9,fontweight="bold")

    arrows = [
        (2.2,4.1,2.8,4.1),(4.6,4.1,5.2,4.1),(7.2,4.1,7.8,4.1),
        (8.7,3.5,8.7,2.7),(8.7,2.5,4.7,4.1),
    ]
    for (x1,y1,x2,y2) in arrows:
        ax.annotate("",xy=(x2,y2),xytext=(x1,y1),
                    arrowprops=dict(arrowstyle="->",color="#ffd54f",lw=2))

    ax.text(5.8,2.5,"Feedback Loop\n(Position Error)",ha="center",color="#ffd54f",fontsize=8,style="italic")
    ax.text(4.65,4.35,"Error (e)",ha="center",color="#90caf9",fontsize=8)
    ax.text(6.2,4.35,"ax,ay",ha="center",color="#90caf9",fontsize=8)

    ax.text(1.2,2.8,"ax = Kp·err_x + Kd·(rabbit_vx - vx)",color="#a5d6a7",fontsize=8,fontfamily="monospace")
    ax.text(1.2,2.3,"ay = Kp·err_y + Kd·(rabbit_vy - vy)",color="#a5d6a7",fontsize=8,fontfamily="monospace")

    path = os.path.join(BASE, "chart_pd_controller.png")
    plt.tight_layout(); plt.savefig(path,dpi=150,bbox_inches="tight",facecolor="#0d1117"); plt.close()
    return path

# ── Generate Deadlock Prevention Diagram ─────────────────────────────────────
def make_deadlock_diagram():
    fig, ax = plt.subplots(figsize=(9,5), facecolor="#0d1117")
    ax.set_facecolor("#0d1117"); ax.set_xlim(0,10); ax.set_ylim(0,6); ax.axis("off")
    ax.set_title("Cell-Based Deadlock Prevention Flow", color="white", fontsize=13, fontweight="bold")

    steps = [
        (0.5,4.5,2.0,1.0,"#1565c0","Drone wants to\nadvance to next cell"),
        (3.0,4.5,2.0,1.0,"#6a1b9a","Check cell registry:\nIs cell claimed?"),
        (5.5,4.5,2.0,1.0,"#00695c","Advance to cell\n✓ Normal flight"),
        (3.0,2.5,2.0,1.0,"#b71c1c","Wait / Hover in place\nStart timer"),
        (5.5,2.5,2.0,1.0,"#e65100","Timer > 2.5s?\nDEADLOCK"),
        (7.5,2.5,2.0,1.0,"#1b5e20","BFS Reroute\naround blocked cells"),
    ]
    for (x,y,w,h,col,txt) in steps:
        rect = FancyBboxPatch((x,y),w,h,boxstyle="round,pad=0.1",
                              facecolor=col,edgecolor="white",linewidth=1.5,alpha=0.9)
        ax.add_patch(rect)
        ax.text(x+w/2,y+h/2,txt,ha="center",va="center",color="white",fontsize=8.5,fontweight="bold")

    ax.annotate("",xy=(3.0,5.0),xytext=(2.5,5.0),arrowprops=dict(arrowstyle="->",color="#ffd54f",lw=2))
    ax.text(3.3,5.35,"No",color="#2ecc71",fontsize=9,fontweight="bold")
    ax.annotate("",xy=(5.5,5.0),xytext=(5.0,5.0),arrowprops=dict(arrowstyle="->",color="#ffd54f",lw=2))
    ax.text(5.05,5.35,"Yes",color="#e74c3c",fontsize=9,fontweight="bold")
    ax.annotate("",xy=(4.0,3.5),xytext=(4.0,4.5),arrowprops=dict(arrowstyle="->",color="#ffd54f",lw=2))
    ax.annotate("",xy=(5.5,3.0),xytext=(5.0,3.0),arrowprops=dict(arrowstyle="->",color="#ffd54f",lw=2))
    ax.annotate("",xy=(7.5,3.0),xytext=(7.5,3.0),arrowprops=dict(arrowstyle="->",color="#ffd54f",lw=2))
    ax.text(3.3,4.1,"Blocked",color="#e74c3c",fontsize=8)
    ax.text(5.05,3.35,"Yes",color="#e74c3c",fontsize=9,fontweight="bold")

    path = os.path.join(BASE, "chart_deadlock.png")
    plt.tight_layout(); plt.savefig(path,dpi=150,bbox_inches="tight",facecolor="#0d1117"); plt.close()
    return path

# ── Generate Drone Speed Comparison Chart ────────────────────────────────────
def make_speed_chart():
    fig, ax = plt.subplots(figsize=(9,4), facecolor="#0d1117")
    ax.set_facecolor("#161b22")
    robots = [f"R{i}" for i in range(1,11)]
    speeds = [0.8,0.8,0.8,0.8,0.8,0.8,0.6,0.8,0.8,0.8]
    colors = ["#e74c3c","#3498db","#2ecc71","#9b59b6","#e67e22",
              "#00bcd4","#e91e8c","#f1c40f","#009688","#ecf0f1"]
    bars = ax.bar(robots, speeds, color=colors, edgecolor="#0d1117", linewidth=1.5, width=0.6)
    ax.set_ylim(0,1.1); ax.set_xlabel("Drone ID", color="white", fontsize=11)
    ax.set_ylabel("Speed (m/s)", color="white", fontsize=11)
    ax.set_title("Drone Speed Configuration", color="white", fontsize=13, fontweight="bold")
    ax.tick_params(colors="white", labelsize=10)
    ax.spines[["top","right"]].set_visible(False)
    ax.spines[["left","bottom"]].set_color("#444")
    ax.yaxis.grid(True, color="#333", linestyle="--", alpha=0.6); ax.set_axisbelow(True)
    for bar,sp in zip(bars,speeds):
        ax.text(bar.get_x()+bar.get_width()/2,bar.get_height()+0.02,f"{sp}",
                ha="center",va="bottom",color="white",fontsize=9,fontweight="bold")
    path = os.path.join(BASE, "chart_speeds.png")
    plt.tight_layout(); plt.savefig(path,dpi=150,bbox_inches="tight",facecolor="#0d1117"); plt.close()
    return path

pd_img       = make_pd_diagram()
deadlock_img = make_deadlock_diagram()
speed_img    = make_speed_chart()
print(f"Charts: {pd_img}, {deadlock_img}, {speed_img}")

# ── Word helpers ──────────────────────────────────────────────────────────────
def set_cell_bg(cell, hx):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
    shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),hx); tcPr.append(shd)

def add_heading(doc,text,level=1,color="1E3A5F"):
    h=doc.add_heading(text,level=level); h.alignment=WD_ALIGN_PARAGRAPH.LEFT
    for r in h.runs: r.font.color.rgb=RGBColor.from_string(color); r.font.bold=True

def add_para(doc,text,size=11,bold=False,color=None,mono=False):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.LEFT
    r=p.add_run(text); r.font.size=Pt(size); r.font.bold=bold
    if mono: r.font.name="Courier New"
    if color: r.font.color.rgb=RGBColor.from_string(color)

def add_code(doc,text):
    p=doc.add_paragraph(); p.style=doc.styles["No Spacing"]
    r=p.add_run(text); r.font.name="Courier New"; r.font.size=Pt(8.5)
    r.font.color.rgb=RGBColor(0x1A,0x5C,0x2A)

def add_img(doc,path,width=Inches(5.5),caption=None):
    if path and os.path.exists(path):
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path,width=width)
        if caption:
            cp=doc.add_paragraph(caption); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in cp.runs: r.font.size=Pt(9); r.font.italic=True; r.font.color.rgb=RGBColor(100,100,100)
    else: doc.add_paragraph(f"[Image not found: {path}]")

def add_table(doc,headers,rows,hbg="1E3A5F",alt="EAF0FB"):
    t=doc.add_table(rows=1+len(rows),cols=len(headers))
    t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    hr=t.rows[0]
    for i,h in enumerate(headers):
        c=hr.cells[i]; c.text=h; set_cell_bg(c,hbg)
        for p in c.paragraphs:
            for r in p.runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for ri,rd in enumerate(rows):
        row=t.rows[ri+1]; bg=alt if ri%2==0 else "FFFFFF"
        for ci,val in enumerate(rd):
            c=row.cells[ci]; c.text=val; set_cell_bg(c,bg)
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)

# ── Load & append ─────────────────────────────────────────────────────────────
doc = Document(OUT_DOCX)

# ── SECTION 6: ALGORITHMS ─────────────────────────────────────────────────────
add_heading(doc,"6. Algorithms in Detail",level=1)

# 6.1 BFS
add_heading(doc,"6.1  BFS Pathfinding",level=2,color="2E5FA0")
add_para(doc,(
    "Breadth-First Search is used in three distinct scenarios: "
    "(A) initial path stitching via resolve_path(), "
    "(B) deadlock rerouting via find_path_excluding(), and "
    "(C) dynamic obstacle rerouting when /dynamic_obstacles is received."
))
add_img(doc,BFS_IMG,width=Inches(5.5),
        caption="Figure 7 — BFS pathfinding flowchart and grid traversal example")
doc.add_paragraph("")
add_para(doc,"Core BFS Implementation:",bold=True)
add_code(doc,
"""visited = {start}
queue   = deque([(start, [start])])
while queue:
    cur, path = queue.popleft()
    for nb in self._adj[cur]:        # O(1) adjacency lookup
        if nb == end:
            return path + [nb]       # Shortest path found
        if nb not in visited:
            visited.add(nb)
            queue.append((nb, path + [nb]))
return None                          # No path exists""")
add_para(doc,"Complexity: O(V+E) where V=80 nodes, E≈120 edges. Guarantees shortest hop-count path.",
         size=10, color="555555")
doc.add_paragraph("")

# 6.2 PD Controller
add_heading(doc,"6.2  PD Controller & Virtual Rabbit Tracking",level=2,color="2E5FA0")
add_para(doc,(
    "Instead of teleporting along grid segments, a Proportional-Derivative controller pursues "
    "a ghost 'virtual rabbit' that moves strictly along the grid path. This produces smooth "
    "inertial motion with natural momentum-based cornering."
))
add_img(doc,pd_img,width=Inches(5.8),
        caption="Figure 8 — PD controller closed-loop block diagram with virtual rabbit tracking")
doc.add_paragraph("")
add_table(doc,["Parameter","Value","Role"],[
    ("Kp (Proportional)","3.5","Closes the position gap between drone and rabbit"),
    ("Kd (Derivative)","2.0","Damps oscillation by matching rabbit's velocity"),
    ("Speed cap","speed × 1.5","Prevents corner runaway / overshoot"),
    ("Timer rate","50 Hz (0.02s)","Physics integration frequency"),
])
doc.add_paragraph("")
add_para(doc,"Step-by-step control law:",bold=True)
add_code(doc,
"""# Step 1: Advance virtual rabbit along path
alpha     += (speed * dt) / edge_length
rabbit_x,y = gm.interpolate(cur_node, next_node, alpha)

# Step 2: Compute errors
err_x = rabbit_x - self.x;  err_y = rabbit_y - self.y
rabbit_vx = (dx/edge_len)*speed;  rabbit_vy = (dy/edge_len)*speed

# Step 3: PD control
ax = Kp*err_x + Kd*(rabbit_vx - vx)
ay = Kp*err_y + Kd*(rabbit_vy - vy)

# Step 4: Euler integration
vx += ax*dt;  x += vx*dt
vy += ay*dt;  y += vy*dt""")
doc.add_paragraph("")

# 6.3 Tilt Model
add_heading(doc,"6.3  Quadcopter Tilt Model",level=2,color="2E5FA0")
add_para(doc,(
    "A real quadcopter generates horizontal force by tilting its thrust vector. "
    "The simulation reverses this: the required acceleration determines the tilt angle. "
    "A spring-damper smooths the transition, giving ~0.3s lag matching real drone footage."
))
add_code(doc,
"""# Rotate world-frame acceleration into body frame
a_fwd = ax*cos(psi) + ay*sin(psi)
a_lat = -ax*sin(psi) + ay*cos(psi)

# Convert to tilt angles (amplified ×1.5 for visual clarity)
target_pitch = atan2(a_fwd, g) * 1.5
target_roll  = -atan2(a_lat, g) * 1.5

# Spring-damper smoothing (spring=40.0, damp=8.0)
pitch_vel += ((target_pitch - pitch)*40.0 - pitch_vel*8.0) * dt
pitch     += pitch_vel * dt""")
doc.add_paragraph("")

# 6.4 Yaw Spring-Damper
add_heading(doc,"6.4  Yaw Spring-Damper",level=2,color="2E5FA0")
add_para(doc,(
    "The drone's heading tracks its true velocity direction, not the grid segment. "
    "A spring-damper with angle-wrapping ensures the drone always rotates via the "
    "shortest arc, with a max yaw rate of 3 rad/s."
))
add_code(doc,
"""if cur_speed > 0.1:
    target_yaw = atan2(vy, vx)
yaw_diff  = angle_diff(target_yaw, current_yaw)   # normalised to (-pi, pi]
yaw_accel = yaw_diff*15.0 - yaw_vel*5.0
yaw_vel   = clamp(yaw_vel + yaw_accel*dt, -3.0, +3.0)
current_yaw += yaw_vel * dt""")
doc.add_paragraph("")

# 6.5 Deadlock Prevention
add_heading(doc,"6.5  Cell-Based Deadlock Prevention",level=2,color="2E5FA0")
add_para(doc,(
    "Each drone broadcasts its claimed cells to /robot_cell_registry. Before advancing, "
    "it checks if the next cell is claimed by another drone. If blocked for >2.5 seconds, "
    "BFS rerouting is triggered around all other drones' cells."
))
add_img(doc,deadlock_img,width=Inches(5.5),
        caption="Figure 9 — Deadlock prevention flowchart: cell check → wait → BFS reroute")
doc.add_paragraph("")

# 6.6 Distance-Priority Yielding
add_heading(doc,"6.6  Distance-Priority Yielding",level=2,color="2E5FA0")
add_para(doc,(
    "A reactive secondary layer for cases between cell boundaries. "
    "Lower robot ID = higher priority. If a drone comes within 2.5 m of a higher-priority "
    "drone, it pauses all PD advancement until the gap widens."
))
add_code(doc,
"""for other_id, other_pos in self._poses_of_others.items():
    dist3d = sqrt((other_pos.x-x)**2 + (other_pos.y-y)**2 + (other_pos.z-z)**2)
    if dist3d < 2.5 and get_id_num(other_id) < get_id_num(self.robot_id):
        self._is_yielding = True   # This drone pauses
        break""")
doc.add_paragraph("")

# 6.7 Dynamic Obstacles
add_heading(doc,"6.7  Dynamic Obstacle Spawning",level=2,color="2E5FA0")
add_para(doc,(
    "graph_visualizer fires a timer at the user-defined spawn_interval. It samples 15 random "
    "unoccupied cells and broadcasts them via /dynamic_obstacles. All robot_nodes receive this "
    "within milliseconds, rebuild their adjacency graph, and BFS-reroute on the next physics tick."
))
add_table(doc,["spawn_interval","Behaviour"],[
    ("0","Timer never created. Grid stays static with only original 12 obstacles."),
    ("1–5s","Very aggressive reshuffling — drones reroute frequently"),
    ("8s (default)","Balanced: periodic obstacle changes without constant disruption"),
    ("30s+","Infrequent reshuffling — nearly static environment"),
])
doc.add_page_break()

# ── SECTION 7: FLIGHT STATE MACHINE ──────────────────────────────────────────
add_heading(doc,"7. Flight State Machine",level=1)
add_para(doc,(
    "Each drone's lifecycle is governed by a four-state FSM defined by the FlightState enum. "
    "The drone begins LANDED, takes off when a valid path is assigned, enters FLYING at cruise "
    "altitude, then lands upon reaching its final destination."
))
add_img(doc,FSM_IMG,width=Inches(5.8),
        caption="Figure 10 — Flight State Machine: LANDED → TAKING_OFF → FLYING → LANDING")
doc.add_paragraph("")
add_table(doc,["State","Z Behaviour","Propeller","Horizontal Movement"],[
    ("LANDED","Fixed at z = 0.0","Stopped","None"),
    ("TAKING_OFF","z increases at 0.6 m/s","+15 rad/s","None — waiting for cruise altitude"),
    ("FLYING","z = 1.2m + hover bob","+20 rad/s","Full PD physics + all avoidance"),
    ("LANDING","z decreases at 0.6 m/s","+10 rad/s","None — tilt resets toward 0"),
])
doc.add_paragraph("")
add_para(doc,"Hover Bobbing (all non-LANDED states):",bold=True)
add_code(doc,"hover_drift_z = 0.05 * sin(hover_time * 2.5)  # ±5 cm at ~0.4 Hz")
add_para(doc,(
    "This sinusoidal altitude oscillation makes flying drones appear alive and distinguishes "
    "hovering-in-place from the LANDED state."
),size=10,color="555555")
doc.add_paragraph("")
add_img(doc,speed_img,width=Inches(5.5),
        caption="Figure 11 — Drone speed configuration: robot7 runs at 0.6 m/s (perimeter loop)")
doc.add_page_break()

doc.save(OUT_DOCX)
print(f"✅ Part 3 SAVED → {OUT_DOCX}")
print("   Sections added: 6 (Algorithms), 7 (Flight State Machine)")
