"""
Part 1: Install deps, generate matplotlib charts, build Word doc (Cover + Overview)
"""
import subprocess, sys, os, glob

# Bootstrap pip if missing
try:
    import pip as _pip_check
except ImportError:
    print("pip not found — installing via apt...")
    subprocess.check_call(["sudo", "apt-get", "install", "-y", "python3-pip"], check=False)
    subprocess.check_call([sys.executable, "-m", "ensurepip", "--upgrade"])
    subprocess.check_call([sys.executable, "-m", "pip", "install", "--upgrade", "pip", "-q"])

def pip_install(pkg):
    subprocess.check_call([sys.executable, "-m", "pip", "install", pkg, "-q", "--user"])

for p in ["python-docx", "matplotlib", "numpy", "Pillow"]:
    pip_install(p)

# Add user site-packages so freshly installed modules are importable
import site
user_site = site.getusersitepackages()
if user_site not in sys.path:
    sys.path.insert(0, user_site)

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE     = r"/home/jatin/ros2_assign_2_10_robo_modified_tonew/turtle_to_quadcopter final with modifications"
IMG_BASE = r"C:\Users\HP\.gemini\antigravity\brain\52bd95c7-02fd-43b7-aea1-8cf37fce8c85"
OUT_DOCX = os.path.join(BASE, "Quadcopter_MultiAgent_Report.docx")
CHART_DIR = BASE  # save charts alongside the docx

def find_img(prefix):
    hits = glob.glob(os.path.join(IMG_BASE, f"{prefix}_*.png"))
    return hits[0] if hits else None

COVER_IMG = find_img("cover_illustration")
GRID_IMG  = find_img("grid_map_diagram")
DRONE_IMG = find_img("drone_architecture_diagram")
ROS_IMG   = find_img("ros_architecture_diagram")
FSM_IMG   = find_img("flight_state_machine")
BFS_IMG   = find_img("bfs_flowchart")

print(f"COVER : {COVER_IMG}")
print(f"GRID  : {GRID_IMG}")
print(f"DRONE : {DRONE_IMG}")
print(f"ROS   : {ROS_IMG}")
print(f"FSM   : {FSM_IMG}")
print(f"BFS   : {BFS_IMG}")

# ── Chart 1: Grid Statistics Pie Chart ────────────────────────────────────────
def chart_grid_pie():
    fig, ax = plt.subplots(figsize=(6,5), facecolor="#0d1117")
    ax.set_facecolor("#0d1117")
    sizes  = [12, 15, 53]
    labels = ["Static Obstacles (12)", "Max Dynamic Obstacles (15)", "Walkable Cells (53 min)"]
    colors = ["#e74c3c", "#e67e22", "#2ecc71"]
    wedges, texts, autotexts = ax.pie(
        sizes, labels=labels, colors=colors, explode=(0.05,0.05,0.05),
        autopct="%1.0f%%", startangle=140,
        textprops={"color":"white","fontsize":10,"fontweight":"bold"},
        wedgeprops={"linewidth":2,"edgecolor":"#0d1117"}
    )
    for at in autotexts:
        at.set_fontsize(10); at.set_color("white")
    ax.set_title("Grid Cell Distribution (80 Total Cells)", color="white", fontsize=13, fontweight="bold", pad=15)
    path = os.path.join(CHART_DIR, "chart_grid_pie.png")
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close()
    return path

# ── Chart 2: Drone Speed Bar Chart ────────────────────────────────────────────
def chart_drone_speeds():
    fig, ax = plt.subplots(figsize=(9,5), facecolor="#0d1117")
    ax.set_facecolor("#161b22")
    robots = [f"robot{i}" for i in range(1,11)]
    speeds = [0.8,0.8,0.8,0.8,0.8,0.8,0.6,0.8,0.8,0.8]
    colors = ["#e74c3c","#3498db","#2ecc71","#9b59b6","#e67e22",
              "#00bcd4","#e91e8c","#f1c40f","#009688","#ecf0f1"]
    bars = ax.bar(robots, speeds, color=colors, edgecolor="#0d1117", linewidth=1.5, width=0.6)
    ax.set_ylim(0,1.1)
    ax.set_xlabel("Drone ID", color="white", fontsize=11)
    ax.set_ylabel("Speed (m/s)", color="white", fontsize=11)
    ax.set_title("Drone Speed Configuration by Robot", color="white", fontsize=13, fontweight="bold")
    ax.tick_params(colors="white", labelsize=9)
    ax.spines[["top","right"]].set_visible(False)
    ax.spines[["left","bottom"]].set_color("#444")
    ax.yaxis.grid(True, color="#333", linestyle="--", alpha=0.6)
    ax.set_axisbelow(True)
    for bar,sp in zip(bars,speeds):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.02,
                f"{sp}", ha="center", va="bottom", color="white", fontsize=9, fontweight="bold")
    path = os.path.join(CHART_DIR, "chart_drone_speeds.png")
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close()
    return path

# ── Chart 3: Timer Architecture ───────────────────────────────────────────────
def chart_timer_arch():
    fig, ax = plt.subplots(figsize=(8,4), facecolor="#0d1117")
    ax.set_facecolor("#0d1117")
    ax.set_xlim(0,10); ax.set_ylim(0,6); ax.axis("off")
    ax.set_title("Timer Architecture per Robot Node", color="white", fontsize=13, fontweight="bold")
    for x, col, title, sub in [
        (0.8, "#1565c0", "Physics Loop\n50 Hz (0.02s)", "PD Control / Collision\nAvoidance / BFS"),
        (5.5, "#00695c", "Marker Refresh\n3 Hz (0.33s)", "RViz Markers\nDrone Visuals"),
    ]:
        rect = FancyBboxPatch((x,3.2), 3.2,1.8, boxstyle="round,pad=0.15",
                              facecolor=col, edgecolor="white", linewidth=1.5, alpha=0.85)
        ax.add_patch(rect)
        ax.text(x+1.6, 4.5, title, ha="center", va="center", color="white", fontsize=10, fontweight="bold")
        ax.text(x+1.6, 3.7, sub,   ha="center", va="center", color="#b0bec5", fontsize=8)
    ax.annotate("", xy=(5.4,4.2), xytext=(4.1,4.2),
                arrowprops=dict(arrowstyle="->", color="#ffd54f", lw=2))
    ax.text(4.75, 4.5, "Parallel\nTimers", ha="center", color="#ffd54f", fontsize=8)
    rect2 = FancyBboxPatch((3.5,0.8), 3.0,1.4, boxstyle="round,pad=0.15",
                           facecolor="#37474f", edgecolor="#90a4ae", linewidth=1.5)
    ax.add_patch(rect2)
    ax.text(5.0,1.5, "robot_node.py\n(ROS 2 Node)", ha="center", va="center",
            color="white", fontsize=10, fontweight="bold")
    ax.annotate("", xy=(2.4,3.1), xytext=(4.2,2.3),
                arrowprops=dict(arrowstyle="->", color="#90a4ae", lw=1.5))
    ax.annotate("", xy=(7.1,3.1), xytext=(5.8,2.3),
                arrowprops=dict(arrowstyle="->", color="#90a4ae", lw=1.5))
    path = os.path.join(CHART_DIR, "chart_timer_arch.png")
    plt.tight_layout()
    plt.savefig(path, dpi=150, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close()
    return path

pie_path   = chart_grid_pie()
speed_path = chart_drone_speeds()
timer_path = chart_timer_arch()
print(f"Charts saved: {pie_path}, {speed_path}, {timer_path}")

# ── Word helpers ───────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color="1E3A5F"):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color); run.font.bold = True
    return h

def add_para(doc, text, bold=False, size=11, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(); p.alignment = align
    run = p.add_run(text); run.font.size = Pt(size); run.font.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)
    return p

def add_img(doc, path, width=Inches(5.5), caption=None):
    if path and os.path.exists(path):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, width=width)
        if caption:
            cp = doc.add_paragraph(caption); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cp.runs:
                r.font.size = Pt(9); r.font.italic = True
                r.font.color.rgb = RGBColor(100,100,100)
    else:
        doc.add_paragraph(f"[Image not found: {path}]")

def add_table(doc, headers, rows, hbg="1E3A5F", alt="EAF0FB"):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = t.rows[0]
    for i,h in enumerate(headers):
        c = hr.cells[i]; c.text = h; set_cell_bg(c, hbg)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri,rd in enumerate(rows):
        row = t.rows[ri+1]; bg = alt if ri%2==0 else "FFFFFF"
        for ci,val in enumerate(rd):
            c = row.cells[ci]; c.text = val; set_cell_bg(c, bg)
            for p in c.paragraphs:
                for r in p.runs: r.font.size = Pt(9)
    return t

# ── Build Document ─────────────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sec.page_width=Cm(21); sec.page_height=Cm(29.7)
sec.left_margin=sec.right_margin=Cm(2.5)
sec.top_margin=sec.bottom_margin=Cm(2.0)
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ── COVER PAGE ────────────────────────────────────────────────────────────────
add_img(doc, COVER_IMG, width=Inches(6.0))
doc.add_paragraph("")
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Quadcopter Multi-Agent Simulation")
r.font.size=Pt(26); r.font.bold=True; r.font.color.rgb=RGBColor(0x1E,0x3A,0x5F)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Complete Technical Report  |  ROS 2 Humble  |  Python 3.10")
r2.font.size=Pt(13); r2.font.italic=True; r2.font.color.rgb=RGBColor(0x55,0x66,0x88)
doc.add_paragraph("")
for line in ["Package: robot_proximity","ROS 2: Humble Hawksbill","Language: Python 3.10","Maintainer: Jatin","License: Apache-2.0"]:
    mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = mp.add_run(line); mr.font.size=Pt(11); mr.font.color.rgb=RGBColor(0x33,0x44,0x55)
doc.add_page_break()

# ── TABLE OF CONTENTS ─────────────────────────────────────────────────────────
add_heading(doc, "Table of Contents", level=1)
toc = [
    ("1","Project Overview","3"), ("2","Repository Structure","4"),
    ("3","ROS 2 Package Configuration","5"), ("4","System Architecture","6"),
    ("5","Component Deep Dive","7"), ("  5.1","graph_manager.py","7"),
    ("  5.2","robot_node.py","8"), ("  5.3","graph_visualizer.py","9"),
    ("  5.4","proximity_monitor.py","10"), ("  5.5","interactive_runner.py","10"),
    ("6","Algorithms in Detail","11"), ("  6.1","BFS Pathfinding","11"),
    ("  6.2","PD Controller & Virtual Rabbit","12"), ("  6.3","Quadcopter Tilt Model","13"),
    ("  6.4","Yaw Spring-Damper","13"), ("  6.5","Cell-Based Deadlock Prevention","14"),
    ("  6.6","Distance-Priority Yielding","14"), ("  6.7","Dynamic Obstacle Spawning","15"),
    ("7","Flight State Machine","16"), ("8","ROS Topics, Services & Parameters","17"),
    ("9","Running the Project — All Modes","18"), ("10","Build Instructions","19"),
    ("11","Grid Reference","20"),
]
add_table(doc, ["§","Section Title","Page"], toc)
doc.add_page_break()

# ── SECTION 1: PROJECT OVERVIEW ───────────────────────────────────────────────
add_heading(doc, "1. Project Overview", level=1)
add_para(doc, (
    "The Quadcopter Multi-Agent Simulation is a fully decentralized, ROS 2-based multi-robot "
    "coordination system simulating up to 10 autonomous quadcopters flying simultaneously in a "
    "shared 3D airspace above a structured grid map. Every drone operates as its own independent "
    "ROS 2 node — there is no central controller. All coordination flows through the ROS 2 "
    "publish/subscribe mesh and service layer."
))
doc.add_paragraph("")
add_table(doc, ["Feature","Description"], [
    ("Realistic Physics","PD-controller pursuit of virtual rabbit produces smooth inertial motion"),
    ("True Quadcopter Dynamics","Pitch and Roll derived mathematically from acceleration vector"),
    ("Rounded Cornering","Momentum causes sweeping arcs at 90-degree turns"),
    ("BFS Pathfinding","Automatic rerouting around static and dynamic obstacles"),
    ("Cell Deadlock Prevention","Grid-cell reservation prevents head-on collisions"),
    ("Priority Yielding","Distance-based yielding; lower ID drones have right-of-way"),
    ("Dynamic Obstacles","User-defined interval controls obstacle layout reshuffle"),
    ("RViz Visualization","Full 3D render: spinning propellers, tilting body, labels"),
    ("Multi-Mode Launch","AUTO mode (pre-stored paths) or MANUAL mode (waypoints)"),
])
doc.add_paragraph("")
add_img(doc, COVER_IMG, width=Inches(5.8),
        caption="Figure 1 — 10 autonomous quadcopters in the simulation environment")
doc.add_page_break()

# ── SECTION 2: REPOSITORY STRUCTURE ──────────────────────────────────────────
add_heading(doc, "2. Repository Structure", level=1)
add_para(doc, "The workspace follows a standard colcon/ament_python layout:")
p_struct = doc.add_paragraph()
p_struct.style = doc.styles["No Spacing"]
rr = p_struct.add_run(
"""turtle_to_quadcopter final with modifications/
│
├── README_A_TO_Z.md
├── src/
│   └── robot_proximity/
│       ├── package.xml
│       ├── setup.py
│       ├── setup.cfg
│       ├── resource/robot_proximity
│       ├── launch/simulation.launch.py
│       ├── config/simulation.rviz
│       └── robot_proximity/
│           ├── __init__.py
│           ├── graph_manager.py
│           ├── robot_node.py
│           ├── graph_visualizer.py
│           ├── proximity_monitor.py
│           ├── interactive_runner.py
│           └── diag_test.py
├── build/   install/   log/
""")
rr.font.name="Courier New"; rr.font.size=Pt(9); rr.font.color.rgb=RGBColor(0x1A,0x5C,0x2A)
doc.add_paragraph("")
add_table(doc, ["File","Purpose"], [
    ("graph_manager.py","10×8 grid, edge construction, BFS pathfinding engine"),
    ("robot_node.py","Core physics agent: PD control, FSM, collision avoidance"),
    ("graph_visualizer.py","RViz grid renderer + dynamic obstacle spawner"),
    ("proximity_monitor.py","Inter-drone distance logger + service trigger"),
    ("interactive_runner.py","Terminal UI orchestrator + launch file generator"),
    ("diag_test.py","Offline path validator (no ROS required)"),
    ("simulation.launch.py","Static 4-robot quick-launch configuration"),
    ("simulation.rviz","RViz2 display panel and camera configuration"),
])
doc.add_page_break()

doc.save(OUT_DOCX)
print(f"\n✅ Part 1 SAVED → {OUT_DOCX}")
